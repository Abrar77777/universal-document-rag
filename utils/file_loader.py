import csv
import json
from pathlib import Path

import docx
import pandas as pd
from pypdf import PdfReader

def load_file(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
        return text

    elif suffix in {".txt", ".md"}:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif suffix == ".docx":
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(paragraphs + table_rows)

    elif suffix == ".csv":
        rows = []
        with open(file_path, newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)

    elif suffix in {".xlsx", ".xls"}:
        sheets = pd.read_excel(file_path, sheet_name=None)
        parts = []
        for sheet_name, df in sheets.items():
            parts.append(f"Sheet: {sheet_name}")
            parts.append(df.to_csv(index=False))
        return "\n".join(parts)

    elif suffix == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return json.dumps(data, indent=2, ensure_ascii=True)

    else:
        raise ValueError("Unsupported file type")

